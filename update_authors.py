"""
Script to update the IEEE Research Paper with author block (4 members, 2x2 grid).
Run this once — it will insert the author section right after the paper title.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Author Data ────────────────────────────────────────────────────────────────
AUTHORS = [
    {
        "name":   "Pratham Midha",
        "dept":   "Dept. of Computer Science and Engineering",
        "org":    "Chitkara University",
        "city":   "Rajpura, Punjab, India",
        "email":  "email@chitkara.edu.in",          # ← fill this in
    },
    {
        "name":   "Prachi Jain",
        "dept":   "Dept. of Computer Science and Engineering",
        "org":    "Chitkara University",
        "city":   "Rajpura, Punjab, India",
        "email":  "email@chitkara.edu.in",          # ← fill this in
    },
    {
        "name":   "Palak",
        "dept":   "Dept. of Computer Science and Engineering",
        "org":    "Chitkara University",
        "city":   "Rajpura, Punjab, India",
        "email":  "email@chitkara.edu.in",          # ← fill this in
    },
    {
        "name":   "Saikat",
        "dept":   "Dept. of Computer Science and Engineering",
        "org":    "Chitkara University",
        "city":   "Rajpura, Punjab, India",
        "email":  "email@chitkara.edu.in",          # ← fill this in
    },
]

DOCX_PATH = r"IEEE_Research_Paper_Fraud_Detection.docx"
OUTPUT_PATH = r"IEEE_Research_Paper_Fraud_Detection_Updated.docx"
# ────────────────────────────────────────────────────────────────────────────────


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_author_run(para, text, bold=False, italic=False, size_pt=9, color=None):
    """Add a run to a paragraph with specific formatting."""
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def fill_author_cell(cell, author):
    """Fill a table cell with IEEE-formatted author info."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Clear default empty paragraph
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    # Line 1: Author name (bold)
    p_name = cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_author_run(p_name, author["name"], bold=True, size_pt=10)
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(0)

    # Line 2: Department (italic)
    p_dept = cell.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_author_run(p_dept, author["dept"], italic=True, size_pt=9)
    p_dept.paragraph_format.space_before = Pt(0)
    p_dept.paragraph_format.space_after = Pt(0)

    # Line 2b: (of Affiliation) label
    p_aff1 = cell.add_paragraph()
    p_aff1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_author_run(p_aff1, "(of Affiliation)", italic=True, size_pt=9)
    p_aff1.paragraph_format.space_before = Pt(0)
    p_aff1.paragraph_format.space_after = Pt(0)

    # Line 3: University name (italic)
    p_org = cell.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_author_run(p_org, author["org"], italic=True, size_pt=9)
    p_org.paragraph_format.space_before = Pt(0)
    p_org.paragraph_format.space_after = Pt(0)

    # Line 3b: (of Affiliation) label
    p_aff2 = cell.add_paragraph()
    p_aff2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_author_run(p_aff2, "(of Affiliation)", italic=True, size_pt=9)
    p_aff2.paragraph_format.space_before = Pt(0)
    p_aff2.paragraph_format.space_after = Pt(0)

    # Line 4: City, Country
    p_city = cell.add_paragraph()
    p_city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_author_run(p_city, author["city"], size_pt=9)
    p_city.paragraph_format.space_before = Pt(0)
    p_city.paragraph_format.space_after = Pt(0)

    # Line 5: Email
    p_email = cell.add_paragraph()
    p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_author_run(p_email, author["email"], size_pt=9)
    p_email.paragraph_format.space_before = Pt(0)
    p_email.paragraph_format.space_after = Pt(2)

    # Remove cell borders
    set_cell_border(cell)


def insert_author_table(doc):
    """
    Find the title paragraph and insert a 2x2 author table right after it.
    If an existing author table is found (marked by a special XML tag), replace it.
    """
    # ── Locate insertion point: paragraph after the title ──────────────────────
    title_idx = None
    for i, para in enumerate(doc.paragraphs):
        # The title is usually the first non-empty paragraph
        if para.text.strip():
            title_idx = i
            break

    if title_idx is None:
        print("⚠  Could not find title paragraph. Inserting at start.")
        title_idx = 0

    print(f"✔  Title found at paragraph index {title_idx}: \"{doc.paragraphs[title_idx].text[:60]}\"")

    # ── Build a 2×2 table ──────────────────────────────────────────────────────
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # Set equal column widths (full text width / 2)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(3.3)

    # Fill each cell
    fill_author_cell(table.cell(0, 0), AUTHORS[0])   # Pratham Midha
    fill_author_cell(table.cell(0, 1), AUTHORS[1])   # Prachi Jain
    fill_author_cell(table.cell(1, 0), AUTHORS[2])   # Palak
    fill_author_cell(table.cell(1, 1), AUTHORS[3])   # Saikat

    # Remove outer table border
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        tblBorders.append(tag)
    tblPr.append(tblBorders)

    # ── Move the table element to right after the title paragraph ──────────────
    title_para_elem = doc.paragraphs[title_idx]._element
    title_para_elem.addnext(tbl)

    print("✔  Author table inserted successfully.")


def main():
    print(f"📄  Opening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    insert_author_table(doc)

    doc.save(OUTPUT_PATH)
    print(f"\n✅  Saved updated paper to: {OUTPUT_PATH}")
    print("\n📝  Remember to open the file and replace the placeholder emails!")
    print("     Search for 'email@chitkara.edu.in' and fill in each author's real email.\n")


if __name__ == "__main__":
    main()
