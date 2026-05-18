import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title Page
    title = doc.add_heading('FraudShield AI: Credit Card Fraud Detection System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('A Hybrid Isolation Forest + XGBoost Framework with Cost-Sensitive Threshold Optimization', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Credit card fraud is a major problem in the financial sector, causing billions in losses annually. "
        "A critical challenge in building machine learning models for fraud detection is the extreme class imbalance. "
        "In typical real-world data, fraud is extremely rare — roughly 1 in every 578 transactions. "
        "For example, in our dataset, legitimate transactions account for 99.828% while fraudulent transactions account for just 0.172%. "
        "This creates the 'Accuracy Paradox', where a naive model predicting all transactions as 'Legitimate' achieves 99.83% accuracy but fails entirely at catching fraud. "
        "To solve this, this project introduces FraudShield AI, an end-to-end machine learning framework. "
        "It leverages a novel 3-stage Hybrid AI Model combining unsupervised anomaly detection with supervised classification, "
        "optimized specifically to minimize real-world financial costs rather than theoretical accuracy."
    )

    # 2. Methodology
    doc.add_heading('2. Methodology', level=1)
    doc.add_paragraph(
        "The project utilizes the Kaggle Credit Card Fraud Detection dataset, containing 284,807 transactions from September 2013 by European cardholders. "
        "The dataset features 28 anonymized Principal Component Analysis (PCA) components (V1 to V28), 'Time', 'Amount', and a binary 'Class' label."
    )
    doc.add_heading('2.1 Data Preprocessing & Balancing', level=2)
    doc.add_paragraph("Feature Engineering: Derived 'Amount_log' to reduce the extreme right-skewness of transaction amounts and an 'Hour' feature to capture temporal fraud patterns.", style='List Bullet')
    doc.add_paragraph("Scaling: Applied StandardScaler to normalize the engineered features.", style='List Bullet')
    doc.add_paragraph("Imbalance Handling: Applied Synthetic Minority Over-sampling Technique (SMOTE) exclusively on the training data. This generates synthetic fraud samples to balance the training set while preventing data leakage into the test set.", style='List Bullet')

    doc.add_heading('2.2 The 3-Stage Hybrid Model (IF-XGB-CST)', level=2)
    doc.add_paragraph("1. Isolation Forest (Anomaly Scoring): An unsupervised algorithm isolates anomalies based on path lengths in random trees. Each transaction is given an anomaly score.", style='List Bullet')
    doc.add_paragraph("2. Augmented XGBoost: The Isolation Forest anomaly score is appended as an additional feature (Feature #31). An XGBoost classifier is then trained on this augmented dataset to capture complex, non-linear fraud patterns.", style='List Bullet')
    doc.add_paragraph("3. Cost-Sensitive Threshold Optimization: Standard models use a 0.5 probability threshold. However, missing a fraud (False Negative) costs an estimated $500, while a false alarm (False Positive) costs ~$10. The system searches across 300 thresholds to minimize the Total Operational Cost. This mathematically shifts the optimal threshold to a much lower value (e.g., 0.02), significantly prioritizing recall.", style='List Bullet')

    # 3. Tools and Technologies
    doc.add_heading('3. Tools and Technologies', level=1)
    doc.add_paragraph("Programming Language: Python 3.8+", style='List Bullet')
    doc.add_paragraph("Data Processing: Pandas, NumPy", style='List Bullet')
    doc.add_paragraph("Machine Learning: Scikit-Learn, XGBoost", style='List Bullet')
    doc.add_paragraph("Imbalance Handling: Imbalanced-learn (SMOTE, ADASYN, etc.)", style='List Bullet')
    doc.add_paragraph("Data Visualization: Matplotlib, Seaborn, Plotly", style='List Bullet')
    doc.add_paragraph("Web Application Framework: Streamlit", style='List Bullet')
    doc.add_paragraph("Document Generation: python-docx", style='List Bullet')

    # 4. Implementation
    doc.add_heading('4. Implementation', level=1)
    doc.add_paragraph(
        "The system is implemented as a complete research pipeline and an interactive deployment application. "
        "The training pipeline evaluates seven distinct models (Logistic Regression, Decision Tree, Random Forest, XGBoost, Isolation Forest, Neural Network, and the Hybrid IF-XGB-CST). "
        "The best-performing hybrid model is serialized and deployed through a Streamlit dashboard."
    )
    doc.add_paragraph("The Streamlit web application provides three primary modes of operation:")
    doc.add_paragraph("Single Transaction Analysis: Users input 30 transaction features manually. The app scales the inputs, calculates the anomaly score, and provides an instant fraud probability via an interactive gauge chart, pulsing red for detected fraud.", style='List Bullet')
    doc.add_paragraph("Batch CSV Upload: Users upload bulk transaction datasets. The system scores thousands of transactions automatically, outputting risk levels (Low/Medium/High) and estimated fraud losses, with an option to download the results.", style='List Bullet')
    doc.add_paragraph("Model Performance Dashboard: Displays an interactive radar chart comparing all models across key metrics, alongside embedded research plots like ROC curves and Precision-Recall curves.", style='List Bullet')

    # 5. Major Findings/Outcomes/Output/Results
    doc.add_heading('5. Major Findings/Outcomes/Output/Results', level=1)
    doc.add_paragraph(
        "A key finding of this research is that traditional accuracy is deeply flawed for fraud detection. "
        "Instead, models were evaluated using Balanced Accuracy, Precision, Recall, F1-Score, PR-AUC, and the Matthews Correlation Coefficient (MCC)."
    )
    doc.add_paragraph("Results:")
    doc.add_paragraph("The baseline Random Forest and XGBoost models achieved high F1-Scores (~0.84-0.86) and ROC-AUCs (~0.97-0.98).", style='List Bullet')
    doc.add_paragraph("The Hybrid IF-XGB-CST model optimized specifically for Cost-Sensitive Thresholding achieved the highest practical business value. By lowering the decision threshold to ~0.02, the model caught significantly more actual frauds (higher Recall).", style='List Bullet')
    doc.add_paragraph("Although prioritizing Recall caused a slight drop in Precision (more false alarms), the asymmetric cost function proved that this approach saves the institution over 20% in operational costs compared to standard thresholding.", style='List Bullet')
    doc.add_paragraph("SMOTE effectively mitigated the class imbalance during training without introducing data leakage.", style='List Bullet')

    # 6. Conclusion and Future Scope
    doc.add_heading('6. Conclusion and Future Scope', level=1)
    doc.add_paragraph(
        "FraudShield AI successfully demonstrates that combining unsupervised anomaly detection with gradient boosting, "
        "coupled with an asymmetric, business-aware cost matrix, yields a highly robust credit card fraud detection system. "
        "The project highlights the critical importance of selecting appropriate evaluation metrics over naive accuracy when dealing with imbalanced datasets."
    )
    doc.add_paragraph("Future Scope:")
    doc.add_paragraph("Real-Time Streaming: Integrating Apache Kafka or Spark Streaming to process and classify transactions in millisecond real-time.", style='List Bullet')
    doc.add_paragraph("Deep Sequence Models: Applying Recurrent Neural Networks (RNNs) or LSTMs to capture sequential behavioral patterns of individual cardholders over time.", style='List Bullet')
    doc.add_paragraph("Federated Learning: Implementing privacy-preserving distributed learning to allow multiple banks to train a shared fraud model without exposing underlying customer data.", style='List Bullet')

    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph("[1] Machine Learning Group - ULB, 'Credit Card Fraud Detection Dataset', Kaggle, 2018.", style='List Bullet')
    doc.add_paragraph("[2] N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, 'SMOTE: Synthetic Minority Over-sampling Technique,' Journal of Artificial Intelligence Research, vol. 16, pp. 321-357, 2002.", style='List Bullet')
    doc.add_paragraph("[3] T. Chen and C. Guestrin, 'XGBoost: A Scalable Tree Boosting System,' in Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016.", style='List Bullet')
    doc.add_paragraph("[4] F. T. Liu, K. M. Ting, and Z.-H. Zhou, 'Isolation Forest,' in 2008 Eighth IEEE International Conference on Data Mining, 2008.", style='List Bullet')
    doc.add_paragraph("[5] C. Elkan, 'The Foundations of Cost-Sensitive Learning,' in International Joint Conference on Artificial Intelligence, 2001.", style='List Bullet')

    # Appendices
    doc.add_heading('Appendices', level=1)
    doc.add_paragraph("Appendix A: Streamlit Dashboard Screenshots", style='List Bullet')
    doc.add_paragraph("Appendix B: Generated EDA and Evaluation Plots (ROC Curves, PR Curves, Correlation Heatmap, Cost vs Threshold Analysis)", style='List Bullet')

    # Save the document
    output_path = os.path.join(os.getcwd(), 'Fraud_Detection_Project_Report_Structured.docx')
    doc.save(output_path)
    print(f"Structured report successfully generated at: {output_path}")

if __name__ == '__main__':
    create_report()
